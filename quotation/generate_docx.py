from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x22, 0x22, 0x22)

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_header_row(table, texts, col_widths=None):
    row = table.rows[0]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, '1a1a2e')

def add_data_row(table, texts, bold_col=None, row_idx=0):
    row = table.add_row()
    for i, text in enumerate(texts):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(10)
        if bold_col is not None and i == bold_col:
            run.bold = True
        if i == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if row_idx % 2 == 0:
            set_cell_shading(cell, 'faf8f4')

def add_section_title(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(18)
    p.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    # Add gold underline via bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="8" w:space="1" w:color="c9a84c"/></w:pBdr>')
    pPr.append(pBdr)

# ==================== HEADER ====================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.space_after = Pt(4)
run = title.add_run('QUOTATION')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.space_after = Pt(12)
run = subtitle.add_run('Website Design, Development & Deployment')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# Divider line
divider = doc.add_paragraph()
divider.space_after = Pt(12)
pPr = divider._p.get_or_add_pPr()
pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="1" w:color="1a1a2e"/></w:pBdr>')
pPr.append(pBdr)

# ==================== PROJECT INFO ====================
info_data = [
    ('Client:', 'Moon Group of Industries Ltd.'),
    ('Prepared By:', 'ovisoft'),
    ('Project:', 'Corporate Website & Web Application'),
    ('Reference No:', 'MG-WEB-2026-001'),
    ('Date:', 'July 05, 2026'),
    ('Valid Until:', 'August 05, 2026'),
    ('Live Demo:', 'https://moon-group-website.web.app'),
]

info_table = doc.add_table(rows=len(info_data), cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (label, value) in enumerate(info_data):
    row = info_table.rows[i]
    cell0 = row.cells[0]
    cell1 = row.cells[1]
    cell0.text = ''
    cell1.text = ''
    p0 = cell0.paragraphs[0]
    run0 = p0.add_run(label)
    run0.bold = True
    run0.font.size = Pt(10)
    run0.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p1 = cell1.paragraphs[0]
    run1 = p1.add_run(value)
    run1.font.size = Pt(10)
    if i == 0 or i == 1:
        run1.bold = True
    if i == len(info_data) - 1:
        run1.font.color.rgb = RGBColor(0xc9, 0xa8, 0x4c)
        run1.bold = True

# Remove borders from info table
for row in info_table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="none" w:sz="0" w:space="0"/><w:left w:val="none" w:sz="0" w:space="0"/><w:bottom w:val="none" w:sz="0" w:space="0"/><w:right w:val="none" w:sz="0" w:space="0"/></w:tcBorders>')
        tcPr.append(tcBorders)

doc.add_paragraph()

# ==================== SCOPE OF WORK ====================
add_section_title(doc, 'কাজের বিবরণ (Scope of Work)')

deliverables = [
    ('1', 'Website Design (UI/UX)', 'Premium Dark ও Light থিম সহ ফুল ডিজাইন, অ্যানিমেশন ও রেসপন্সিভ লেআউট (Mobile + Desktop)'),
    ('2', 'Hero Section', 'Full-screen ইমেজ স্লাইডার, প্রপার্টি সার্চ ফিল্টার, CTA বাটন'),
    ('3', 'Property Showcase', 'সকল প্রজেক্টের গ্রিড গ্যালারি, ফিল্টারিং (Residential / Commercial), ডিটেইল পেজ'),
    ('4', 'About Us Page', 'কোম্পানির ইতিহাস, ফাউন্ডার প্রোফাইল, মাইলস্টোন স্ট্যাটিস্টিক্স'),
    ('5', 'Group Concerns Page', '৮+ সিস্টার কনসার্নের শোকেস (Real Estate, Construction, Garments, Hotel ইত্যাদি)'),
    ('6', 'Virtual Blueprint Planner', 'ইন্টারঅ্যাক্টিভ 2D ফ্লোর প্ল্যান — ভিজিটররা রুম সিলেক্ট করে স্পেসিফিকেশন দেখতে পারবে'),
    ('7', 'Visit Booking System', 'সাইট ভিজিটের জন্য অনলাইন বুকিং ফর্ম (তারিখ, সময়, প্রপার্টি সিলেক্ট)'),
    ('8', 'AI Support Chat', 'Gemini AI চ্যাটবট — ভিজিটরদের প্রশ্নের উত্তর দেবে অটোমেটিক্যালি'),
    ('9', 'Land Partnership Portal', 'জমির মালিকদের জন্য জয়েন্ট ভেঞ্চার প্রস্তাব সাবমিশন ফর্ম'),
    ('10', 'Construction Status', 'চলমান প্রজেক্টের নির্মাণ অগ্রগতি ট্র্যাকিং পেজ'),
    ('11', 'NRB Investment Page', 'প্রবাসী বিনিয়োগকারীদের জন্য বিশেষ পেজ'),
    ('12', 'Referral Program', 'রেফারেল কমিশন প্রোগ্রামের তথ্য পেজ'),
    ('13', 'Contact Us Page', 'কন্ট্যাক্ট ফর্ম, ঠিকানা, ম্যাপ, হটলাইন'),
    ('14', 'Customer Testimonials', 'ক্লায়েন্ট রিভিউ স্লাইডার'),
    ('15', 'Admin CMS Dashboard', 'ওয়েবসাইটের সব কনটেন্ট ম্যানেজ করার অ্যাডমিন প্যানেল (Add/Edit/Delete Property, Slides, About ইত্যাদি)'),
    ('16', 'Database Setup', 'Firebase Cloud Firestore ডাটাবেজ — বুকিং, পার্টনারশিপ, প্রপার্টি ডাটা সংরক্ষণ'),
    ('17', 'Domain Purchase', 'কাস্টম ডোমেইন moongroupofind.ltd কিনে কানেক্ট করে দেওয়া হবে'),
    ('18', 'Hosting & Deployment', 'Firebase Hosting-এ লাইভ ডেপ্লয়মেন্ট (SSL সার্টিফিকেট + CDN সহ)'),
]

scope_table = doc.add_table(rows=1, cols=3)
scope_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(scope_table, ['#', 'Item', 'Details'])
for idx, (num, name, desc) in enumerate(deliverables):
    add_data_row(scope_table, [num, name, desc], bold_col=1, row_idx=idx)

# ==================== PRICING ====================
add_section_title(doc, 'মোট খরচ (Total Cost)')

price_p = doc.add_paragraph()
price_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
price_p.space_before = Pt(12)
price_p.space_after = Pt(4)

label_run = price_p.add_run('সর্বমোট প্রজেক্ট খরচ\n')
label_run.font.size = Pt(11)
label_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

amount_run = price_p.add_run('৳ 90,000 BDT')
amount_run.bold = True
amount_run.font.size = Pt(26)
amount_run.font.color.rgb = RGBColor(0xc9, 0xa8, 0x4c)

note_p = doc.add_paragraph()
note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
note_p.space_after = Pt(12)
note_run = note_p.add_run('ডোমেইন (moongroupofind.ltd) ক্রয় + Firebase হোস্টিং + সম্পূর্ণ ওয়েবসাইট ডেভেলপমেন্ট সহ')
note_run.font.size = Pt(9)
note_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ==================== PAYMENT ====================
add_section_title(doc, 'পেমেন্ট পদ্ধতি (Payment Schedule)')

pay_table = doc.add_table(rows=1, cols=3)
pay_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(pay_table, ['কিস্তি', 'কখন দিতে হবে', 'পরিমাণ'])

row1 = pay_table.add_row()
row1.cells[0].text = ''
p = row1.cells[0].paragraphs[0]
run = p.add_run('১ম কিস্তি')
run.bold = True
run.font.size = Pt(10)
row1.cells[1].text = 'চুক্তি স্বাক্ষর ও কাজ শুরুর সময়'
row1.cells[1].paragraphs[0].runs[0].font.size = Pt(10) if row1.cells[1].paragraphs[0].runs else None
p2 = row1.cells[2].paragraphs[0]
p2.text = ''
run2 = p2.add_run('৳ 45,000')
run2.bold = True
run2.font.size = Pt(10)
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

row2 = pay_table.add_row()
row2.cells[0].text = ''
p = row2.cells[0].paragraphs[0]
run = p.add_run('২য় কিস্তি')
run.bold = True
run.font.size = Pt(10)
row2.cells[1].text = 'ফাইনাল ডেলিভারি ও লাইভ ডেপ্লয়মেন্টের পর'
p2 = row2.cells[2].paragraphs[0]
p2.text = ''
run2 = p2.add_run('৳ 45,000')
run2.bold = True
run2.font.size = Pt(10)
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_cell_shading(row2.cells[0], 'faf8f4')
set_cell_shading(row2.cells[1], 'faf8f4')
set_cell_shading(row2.cells[2], 'faf8f4')

total_row = pay_table.add_row()
total_row.cells[0].text = ''
total_row.cells[1].text = ''
total_row.cells[2].text = ''
merge_cell = total_row.cells[0].merge(total_row.cells[1])
p = merge_cell.paragraphs[0]
run = p.add_run('সর্বমোট')
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
p2 = total_row.cells[2].paragraphs[0]
run2 = p2.add_run('৳ 90,000')
run2.bold = True
run2.font.size = Pt(11)
run2.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_cell_shading(merge_cell, 'f5f0e5')
set_cell_shading(total_row.cells[2], 'f5f0e5')

# Payment note
pay_note = doc.add_paragraph()
pay_note.space_before = Pt(8)
run = pay_note.add_run('💳 পেমেন্ট মাধ্যম: ')
run.bold = True
run.font.size = Pt(10)
run = pay_note.add_run('ব্যাংক ট্রান্সফার, bKash, বা Nagad — প্রতিটি পেমেন্টের রশিদ দেওয়া হবে।')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# ==================== WHAT'S INCLUDED ====================
add_section_title(doc, 'এই খরচের মধ্যে যা যা অন্তর্ভুক্ত')

inc_table = doc.add_table(rows=1, cols=3)
inc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(inc_table, ['#', 'বিষয়', 'বিবরণ'])

included_items = [
    ('✅', 'সম্পূর্ণ ওয়েবসাইট', 'উপরে উল্লেখিত ১৮টি আইটেমের সবকিছু'),
    ('✅', 'ডোমেইন ক্রয়', 'moongroupofind.ltd ডোমেইন কিনে কানেক্ট করে দেওয়া হবে'),
    ('✅', 'হোস্টিং', 'Firebase Hosting — SSL Certificate + CDN সহ'),
    ('✅', 'সোর্স কোড', 'পূর্ণ পেমেন্টের পর সোর্স কোডের মালিকানা ক্লায়েন্টকে দেওয়া হবে'),
    ('✅', 'ফ্রি সাপোর্ট', 'ডেলিভারির পর ৩ মাস ফ্রি বাগ ফিক্স ও ছোটখাটো আপডেট'),
]
for idx, (check, name, desc) in enumerate(included_items):
    add_data_row(inc_table, [check, name, desc], bold_col=1, row_idx=idx)

# ==================== TIMELINE ====================
add_section_title(doc, 'সময়সীমা (Timeline)')

tl_p = doc.add_paragraph()
tl_p.space_before = Pt(6)
run = tl_p.add_run('📅 চুক্তি ও ১ম কিস্তি পাওয়ার পর ২৫ কার্যদিবসের মধ্যে সম্পূর্ণ প্রজেক্ট ডেলিভারি ও লাইভ ডেপ্লয়মেন্ট করা হবে।')
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

tl_note = doc.add_paragraph()
run = tl_note.add_run('ক্লায়েন্টের কাছ থেকে সময়মতো কনটেন্ট (ছবি, টেক্সট, প্রজেক্ট তথ্য) পেলে নির্ধারিত সময়ে কাজ শেষ হবে।')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# ==================== TERMS ====================
add_section_title(doc, 'শর্তাবলি')

terms = [
    'এই কোটেশনটি ইস্যুর তারিখ থেকে ৩০ দিন পর্যন্ত কার্যকর।',
    'ক্লায়েন্ট প্রয়োজনীয় কনটেন্ট (টেক্সট, ছবি, লোগো) নির্ধারিত সময়ে সরবরাহ করবেন।',
    'কাজ শুরুর পর স্কোপে কোনো পরিবর্তন হলে সময় ও খরচ পরিবর্তন হতে পারে।',
    'পূর্ণ পেমেন্টের পর সোর্স কোডের মালিকানা ক্লায়েন্টের কাছে হস্তান্তর হবে।',
]
for i, term in enumerate(terms):
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    run = p.add_run(f'{i+1}. {term}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# ==================== FOOTER ====================
doc.add_paragraph()
footer_divider = doc.add_paragraph()
pPr = footer_divider._p.get_or_add_pPr()
pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="8" w:space="1" w:color="1a1a2e"/></w:pBdr>')
pPr.append(pBdr)

thanks = doc.add_paragraph()
thanks.alignment = WD_ALIGN_PARAGRAPH.CENTER
thanks.space_before = Pt(10)
run = thanks.add_run('ধন্যবাদ। আমরা Moon Group of Industries Ltd. এর সাথে কাজ করতে আগ্রহী।')
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

ref = doc.add_paragraph()
ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ref.add_run('Ref: MG-WEB-2026-001 | Prepared By: ovisoft | Date: July 05, 2026')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# Save
output_path = r'e:\OneDrive\WORK\PROJECT\MOON GROUP\moon-group-website\quotation\Quotation_Moon_Group_Website_Ovisoft.docx'
doc.save(output_path)
print(f'DOCX saved: {output_path}')
